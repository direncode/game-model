"""BTUT Dynamic Report Generator — PDF/DOCX reports from BTUT signal + RAG analysis.

Combines BTUT structural data (scores, fingerprint, cluster context) with
RAG-sourced primary evidence (cited filing excerpts, abstracts) into
downloadable reports. Uses existing ReportLab/python-docx infrastructure.
"""

from __future__ import annotations

import io
import json
import logging
from dataclasses import dataclass, field
from datetime import datetime

logger = logging.getLogger(__name__)


@dataclass
class BTUTReportContent:
    """Rendered report content."""
    content_bytes: bytes = b""
    format: str = "pdf"
    filename: str = ""
    report_type: str = "full_analysis"
    generated_at: str = ""
    entity_keys: list[str] = field(default_factory=list)
    dataset_id: str = ""


class BTUTReportGenerator:
    """Generate structured PDF/DOCX/JSON reports from BTUT + RAG data."""

    async def generate(
        self,
        entity_data: list[dict],
        rag_syntheses: list[dict],
        dataset_id: str,
        report_format: str = "pdf",
        report_type: str = "full_analysis",
    ) -> BTUTReportContent:
        """Generate a report combining BTUT data with RAG analysis.

        Args:
            entity_data: List of analyze() results from query engine
            rag_syntheses: List of RAGSynthesis.to_dict() results
            dataset_id: Dataset identifier
            report_format: "pdf", "docx", or "json"
            report_type: "full_analysis" or "executive_summary"
        """
        report_data = self._build_structure(entity_data, rag_syntheses, dataset_id, report_type)

        if report_format == "pdf":
            content = self._render_pdf(report_data)
        elif report_format == "docx":
            content = self._render_docx(report_data)
        else:
            content = json.dumps(report_data, indent=2, default=str).encode("utf-8")

        entity_keys = [e.get("ticker") or e.get("entity_key", "") for e in entity_data]
        filename = f"btut_report_{'_'.join(entity_keys[:3])}_{datetime.utcnow().strftime('%Y%m%d')}.{report_format}"

        return BTUTReportContent(
            content_bytes=content,
            format=report_format,
            filename=filename,
            report_type=report_type,
            generated_at=datetime.utcnow().isoformat(),
            entity_keys=entity_keys,
            dataset_id=dataset_id,
        )

    def _build_structure(
        self, entities: list[dict], syntheses: list[dict],
        dataset_id: str, report_type: str,
    ) -> dict:
        """Build the report data structure."""
        sections = []

        for i, (entity, synthesis) in enumerate(zip(entities, syntheses)):
            scores = entity.get("scores", {})
            lattice = entity.get("lattice_profile", {})
            cluster = entity.get("cluster", {})
            metadata = entity.get("metadata", {})

            section = {
                "entity_name": entity.get("company_name") or entity.get("ticker") or entity.get("entity_key", ""),
                "entity_type": entity.get("entity_type", ""),
                "entity_key": entity.get("ticker") or entity.get("cik", ""),

                "scores": {
                    "composite": scores.get("composite", 0),
                    "diversity": scores.get("diversity", 0),
                    "reconstruction": scores.get("reconstruction", 0),
                    "anomaly": scores.get("anomaly", 0),
                },

                "structural_role": metadata.get("structural_role", {}).get("role", ""),
                "role_description": metadata.get("structural_role", {}).get("description", ""),

                "fingerprint": lattice.get("fingerprint_48bit", ""),
                "flips": lattice.get("total_flips", 0),
                "flip_rate": lattice.get("flip_rate", 0),

                "cluster_id": cluster.get("id", -1),
                "cluster_size": cluster.get("total_members", 0),
                "peers": cluster.get("peers", [])[:5],

                "executive_summary": synthesis.get("executive_summary", ""),
                "structural_analysis": synthesis.get("structural_analysis", ""),
                "anomaly_explanation": synthesis.get("anomaly_explanation", ""),
                "risk_assessment": synthesis.get("risk_assessment", ""),
                "opportunity_assessment": synthesis.get("opportunity_assessment", ""),
                "confidence_note": synthesis.get("confidence_note", ""),
                "citations": synthesis.get("source_evidence", []),
            }
            sections.append(section)

        return {
            "report_type": report_type,
            "dataset_id": dataset_id,
            "generated_at": datetime.utcnow().isoformat(),
            "entity_count": len(sections),
            "sections": sections,
        }

    def _render_pdf(self, data: dict) -> bytes:
        """Render as PDF using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.colors import HexColor
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=50, bottomMargin=50)
            styles = getSampleStyleSheet()

            # Custom styles
            styles.add(ParagraphStyle("SectionHeader", parent=styles["Heading2"],
                                       textColor=HexColor("#00d4ff"), spaceAfter=8))
            styles.add(ParagraphStyle("BodyText2", parent=styles["Normal"],
                                       fontSize=9, leading=12, spaceAfter=6))
            styles.add(ParagraphStyle("Citation", parent=styles["Normal"],
                                       fontSize=8, leading=10, leftIndent=20,
                                       textColor=HexColor("#888888")))

            story = []

            # Title
            story.append(Paragraph("BTUT Intelligence Report", styles["Title"]))
            story.append(Paragraph(
                f"Dataset: {data['dataset_id']} | Generated: {data['generated_at'][:10]} | "
                f"Entities: {data['entity_count']}",
                styles["Normal"],
            ))
            story.append(Spacer(1, 20))

            for section in data["sections"]:
                # Entity header
                story.append(Paragraph(
                    f"{section['entity_name']} [{section['entity_key']}]",
                    styles["Heading1"],
                ))
                story.append(Paragraph(
                    f"Type: {section['entity_type']} | Role: {section['structural_role']} | "
                    f"Cluster: #{section['cluster_id']} ({section['cluster_size']} members)",
                    styles["Normal"],
                ))
                story.append(Spacer(1, 10))

                # Scores table
                score_data = [
                    ["Composite", "Diversity", "Reconstruction", "Anomaly"],
                    [
                        f"{section['scores']['composite']:.4f}",
                        f"{section['scores']['diversity']:.4f}",
                        f"{section['scores']['reconstruction']:.4f}",
                        f"{section['scores']['anomaly']:.4f}",
                    ],
                ]
                t = Table(score_data, colWidths=[120, 120, 120, 120])
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1a1a1a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#00d4ff")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#333333")),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(t)
                story.append(Spacer(1, 12))

                # Executive Summary
                if section.get("executive_summary"):
                    story.append(Paragraph("Executive Summary", styles["SectionHeader"]))
                    story.append(Paragraph(section["executive_summary"][:1500], styles["BodyText2"]))
                    story.append(Spacer(1, 8))

                # Anomaly Explanation
                if section.get("anomaly_explanation"):
                    story.append(Paragraph("Anomaly Explanation", styles["SectionHeader"]))
                    story.append(Paragraph(section["anomaly_explanation"][:1500], styles["BodyText2"]))
                    story.append(Spacer(1, 8))

                # Source Evidence
                if section.get("citations"):
                    story.append(Paragraph("Primary Source Evidence", styles["SectionHeader"]))
                    for cit in section["citations"][:5]:
                        excerpt = cit.get("excerpt", "")[:300]
                        relevance = cit.get("relevance", "")[:200]
                        story.append(Paragraph(f"<b>{excerpt}</b>", styles["Citation"]))
                        if relevance:
                            story.append(Paragraph(f"Relevance: {relevance}", styles["Citation"]))
                        story.append(Spacer(1, 4))

                # Risk + Opportunity
                if section.get("risk_assessment"):
                    story.append(Paragraph("Risk Assessment", styles["SectionHeader"]))
                    story.append(Paragraph(section["risk_assessment"][:1000], styles["BodyText2"]))

                if section.get("opportunity_assessment"):
                    story.append(Paragraph("Opportunity Assessment", styles["SectionHeader"]))
                    story.append(Paragraph(section["opportunity_assessment"][:1000], styles["BodyText2"]))

                # Confidence
                if section.get("confidence_note"):
                    story.append(Paragraph("Confidence Assessment", styles["SectionHeader"]))
                    story.append(Paragraph(section["confidence_note"][:500], styles["BodyText2"]))

                story.append(Spacer(1, 30))

            doc.build(story)
            return buffer.getvalue()

        except ImportError:
            logger.warning("ReportLab not available; returning JSON")
            return json.dumps(data, indent=2, default=str).encode("utf-8")

    def _render_docx(self, data: dict) -> bytes:
        """Render as DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor

            doc = Document()

            # Title
            doc.add_heading("BTUT Intelligence Report", 0)
            doc.add_paragraph(
                f"Dataset: {data['dataset_id']} | Generated: {data['generated_at'][:10]} | "
                f"Entities: {data['entity_count']}"
            )

            for section in data["sections"]:
                doc.add_heading(
                    f"{section['entity_name']} [{section['entity_key']}]", 1
                )
                doc.add_paragraph(
                    f"Type: {section['entity_type']} | Role: {section['structural_role']} | "
                    f"Cluster: #{section['cluster_id']} ({section['cluster_size']} members)"
                )

                # Scores table
                table = doc.add_table(rows=2, cols=4)
                table.style = "Table Grid"
                for i, header in enumerate(["Composite", "Diversity", "Reconstruction", "Anomaly"]):
                    table.rows[0].cells[i].text = header
                    table.rows[1].cells[i].text = f"{section['scores'].get(header.lower(), 0):.4f}"

                # Sections
                for title, key in [
                    ("Executive Summary", "executive_summary"),
                    ("Anomaly Explanation", "anomaly_explanation"),
                    ("Risk Assessment", "risk_assessment"),
                    ("Opportunity Assessment", "opportunity_assessment"),
                    ("Confidence", "confidence_note"),
                ]:
                    content = section.get(key, "")
                    if content:
                        doc.add_heading(title, 2)
                        doc.add_paragraph(content[:2000])

                # Citations
                if section.get("citations"):
                    doc.add_heading("Primary Source Evidence", 2)
                    for cit in section["citations"][:5]:
                        p = doc.add_paragraph()
                        run = p.add_run(cit.get("excerpt", "")[:300])
                        run.italic = True
                        if cit.get("relevance"):
                            doc.add_paragraph(f"Relevance: {cit['relevance'][:200]}")

                doc.add_page_break()

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            logger.warning("python-docx not available; returning JSON")
            return json.dumps(data, indent=2, default=str).encode("utf-8")
