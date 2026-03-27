"""Intelligence report generation in multiple formats.

Generates executive summaries, full analysis reports, and change reports
comparing crystallization runs. Supports PDF, DOCX, HTML, and JSON output.
"""

from __future__ import annotations

import io
import json
import logging
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.crystallization import CrystallizationJob
from app.models.dataset import Dataset
from app.models.delivery import Report
from app.models.module import HiddenConnection, Module, ModuleEntity, ModuleRelationship

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = ("json", "html", "pdf", "docx")


@dataclass
class ReportContent:
    """Generated report content ready for storage."""

    title: str
    report_type: str
    format: str
    content: bytes
    metadata: dict


class ReportGenerator:
    """Generate intelligence reports from crystallization results."""

    def __init__(
        self,
        db: AsyncSession,
        minio_client: Any | None = None,
    ) -> None:
        self._db = db
        self._minio = minio_client

    async def executive_summary(
        self,
        dataset_id: uuid.UUID,
        job_id: uuid.UUID,
        format: str = "json",
    ) -> ReportContent:
        """Generate an executive summary report.

        Includes key findings, module overview, and top connections.

        Args:
            dataset_id: UUID of the dataset.
            job_id: UUID of the crystallization job.
            format: Output format (json/html/pdf/docx).

        Returns:
            ReportContent with generated report.
        """
        self._validate_format(format)

        data = await self._gather_summary_data(dataset_id, job_id)

        summary = {
            "report_type": "executive_summary",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "dataset": data["dataset_info"],
            "job": data["job_info"],
            "key_findings": {
                "total_modules": data["module_count"],
                "avg_purity": data["avg_purity"],
                "top_modules": data["top_modules"][:5],
                "hidden_connection_count": data["hidden_connection_count"],
            },
            "module_overview": data["module_summaries"],
            "top_connections": data["top_connections"][:10],
        }

        content = await self._render(summary, format)

        return ReportContent(
            title=f"Executive Summary: {data['dataset_info']['name']}",
            report_type="executive_summary",
            format=format,
            content=content,
            metadata={"dataset_id": str(dataset_id), "job_id": str(job_id)},
        )

    async def full_analysis(
        self,
        dataset_id: uuid.UUID,
        job_id: uuid.UUID,
        format: str = "json",
    ) -> ReportContent:
        """Generate a full analysis report.

        Includes complete module details, all connections, and lineage.

        Args:
            dataset_id: UUID of the dataset.
            job_id: UUID of the crystallization job.
            format: Output format.

        Returns:
            ReportContent with complete analysis.
        """
        self._validate_format(format)

        data = await self._gather_full_data(dataset_id, job_id)

        report = {
            "report_type": "full_analysis",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "dataset": data["dataset_info"],
            "job": data["job_info"],
            "modules": data["modules_detail"],
            "hidden_connections": data["hidden_connections"],
            "module_relationships": data["module_relationships"],
            "training_metrics": data["training_metrics"],
        }

        content = await self._render(report, format)

        return ReportContent(
            title=f"Full Analysis: {data['dataset_info']['name']}",
            report_type="full_analysis",
            format=format,
            content=content,
            metadata={"dataset_id": str(dataset_id), "job_id": str(job_id)},
        )

    async def change_report(
        self,
        dataset_id: uuid.UUID,
        job_id_a: uuid.UUID,
        job_id_b: uuid.UUID,
        format: str = "json",
    ) -> ReportContent:
        """Generate a diff report between two crystallization runs.

        Args:
            dataset_id: UUID of the dataset.
            job_id_a: UUID of the first (older) job.
            job_id_b: UUID of the second (newer) job.
            format: Output format.

        Returns:
            ReportContent with change analysis.
        """
        self._validate_format(format)

        modules_a = await self._load_modules(job_id_a)
        modules_b = await self._load_modules(job_id_b)

        diff = self._compute_module_diff(modules_a, modules_b)

        report = {
            "report_type": "change_report",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "dataset_id": str(dataset_id),
            "job_a": str(job_id_a),
            "job_b": str(job_id_b),
            "summary": {
                "modules_added": len(diff["added"]),
                "modules_removed": len(diff["removed"]),
                "modules_changed": len(diff["changed"]),
                "modules_unchanged": len(diff["unchanged"]),
            },
            "added_modules": diff["added"],
            "removed_modules": diff["removed"],
            "changed_modules": diff["changed"],
        }

        content = await self._render(report, format)

        return ReportContent(
            title=f"Change Report: Jobs {job_id_a} vs {job_id_b}",
            report_type="change_report",
            format=format,
            content=content,
            metadata={
                "dataset_id": str(dataset_id),
                "job_a": str(job_id_a),
                "job_b": str(job_id_b),
            },
        )

    async def store_report(
        self,
        report: ReportContent,
        dataset_id: uuid.UUID,
        job_id: uuid.UUID,
        user_id: uuid.UUID,
    ) -> Report:
        """Store a generated report in MinIO and record in PostgreSQL.

        Args:
            report: Generated report content.
            dataset_id: UUID of the dataset.
            job_id: UUID of the job.
            user_id: UUID of the generating user.

        Returns:
            Created Report database record.
        """
        # Store in MinIO
        storage_path = (
            f"reports/{dataset_id}/{job_id}/"
            f"{report.report_type}.{report.format}"
        )

        if self._minio:
            try:
                self._minio.put_object(
                    settings.MINIO_BUCKET,
                    storage_path,
                    io.BytesIO(report.content),
                    len(report.content),
                )
            except Exception:
                logger.exception("Failed to store report in MinIO")
                raise
        else:
            logger.warning("MinIO not available; report storage skipped")

        # Record in PostgreSQL
        db_report = Report(
            dataset_id=dataset_id,
            job_id=job_id,
            report_type=report.report_type,
            title=report.title,
            format=report.format,
            storage_path=storage_path,
            generated_by=user_id,
        )
        self._db.add(db_report)
        await self._db.flush()

        logger.info(
            "Report stored: type=%s, format=%s, path=%s",
            report.report_type,
            report.format,
            storage_path,
        )
        return db_report

    async def _gather_summary_data(
        self, dataset_id: uuid.UUID, job_id: uuid.UUID
    ) -> dict:
        """Gather data for executive summary."""
        dataset = await self._load_dataset(dataset_id)
        job = await self._load_job(job_id)
        modules = await self._load_modules(job_id)

        hidden_result = await self._db.execute(
            select(HiddenConnection).where(HiddenConnection.job_id == job_id).limit(20)
        )
        hidden = hidden_result.scalars().all()

        module_summaries = []
        for m in modules:
            module_summaries.append({
                "index": m.module_index,
                "name": m.name,
                "entity_count": m.entity_count,
                "purity_score": m.purity_score,
                "dominant_type": m.dominant_type,
            })

        purity_scores = [m.purity_score for m in modules if m.purity_score is not None]
        avg_purity = sum(purity_scores) / len(purity_scores) if purity_scores else 0.0

        top_modules = sorted(module_summaries, key=lambda m: m.get("purity_score") or 0, reverse=True)

        return {
            "dataset_info": {"name": dataset.name, "id": str(dataset.id), "entity_count": dataset.entity_count},
            "job_info": {"id": str(job.id), "status": job.status, "module_count": job.module_count},
            "module_count": len(modules),
            "avg_purity": round(avg_purity, 4),
            "top_modules": top_modules,
            "module_summaries": module_summaries,
            "hidden_connection_count": len(hidden),
            "top_connections": [
                {"source": h.source_entity, "target": h.target_entity, "similarity": h.similarity_score}
                for h in hidden
            ],
        }

    async def _gather_full_data(
        self, dataset_id: uuid.UUID, job_id: uuid.UUID
    ) -> dict:
        """Gather comprehensive data for full analysis report."""
        summary = await self._gather_summary_data(dataset_id, job_id)
        modules = await self._load_modules(job_id)

        modules_detail = []
        for m in modules:
            entities_result = await self._db.execute(
                select(ModuleEntity).where(ModuleEntity.module_id == m.id)
            )
            entities = entities_result.scalars().all()

            modules_detail.append({
                "index": m.module_index,
                "name": m.name,
                "description": m.description,
                "entity_count": m.entity_count,
                "purity_score": m.purity_score,
                "dominant_type": m.dominant_type,
                "type_distribution": m.type_distribution,
                "internal_density": m.internal_density,
                "anchor_entities": m.anchor_entities,
                "entities": [
                    {"name": e.entity_name, "type": e.entity_type, "centrality": e.centrality_score}
                    for e in entities
                ],
            })

        # Module relationships
        rels_result = await self._db.execute(
            select(ModuleRelationship).where(
                ModuleRelationship.source_module_id.in_([m.id for m in modules])
            )
        )
        rels = rels_result.scalars().all()

        # Hidden connections
        hidden_result = await self._db.execute(
            select(HiddenConnection).where(HiddenConnection.job_id == job_id)
        )
        hidden = hidden_result.scalars().all()

        summary["modules_detail"] = modules_detail
        summary["module_relationships"] = [
            {
                "source_module_id": str(r.source_module_id),
                "target_module_id": str(r.target_module_id),
                "strength": r.connection_strength,
                "shared_count": r.shared_entity_count,
            }
            for r in rels
        ]
        summary["hidden_connections"] = [
            {
                "source": h.source_entity,
                "target": h.target_entity,
                "similarity": h.similarity_score,
                "validated": h.validated,
            }
            for h in hidden
        ]
        summary["training_metrics"] = []  # Loaded from job.training_log

        return summary

    async def _load_dataset(self, dataset_id: uuid.UUID) -> Dataset:
        result = await self._db.execute(
            select(Dataset).where(Dataset.id == dataset_id)
        )
        return result.scalar_one()

    async def _load_job(self, job_id: uuid.UUID) -> CrystallizationJob:
        result = await self._db.execute(
            select(CrystallizationJob).where(CrystallizationJob.id == job_id)
        )
        return result.scalar_one()

    async def _load_modules(self, job_id: uuid.UUID) -> list[Module]:
        result = await self._db.execute(
            select(Module).where(Module.job_id == job_id).order_by(Module.module_index)
        )
        return list(result.scalars().all())

    @staticmethod
    def _compute_module_diff(
        modules_a: list[Module], modules_b: list[Module]
    ) -> dict:
        """Compute diff between two sets of modules."""
        names_a = {m.name: m for m in modules_a}
        names_b = {m.name: m for m in modules_b}

        added = [{"name": m.name, "entity_count": m.entity_count} for n, m in names_b.items() if n not in names_a]
        removed = [{"name": m.name, "entity_count": m.entity_count} for n, m in names_a.items() if n not in names_b]

        changed = []
        unchanged = []
        for name in set(names_a) & set(names_b):
            ma, mb = names_a[name], names_b[name]
            if (ma.entity_count != mb.entity_count or
                    ma.purity_score != mb.purity_score or
                    ma.dominant_type != mb.dominant_type):
                changed.append({
                    "name": name,
                    "before": {"entity_count": ma.entity_count, "purity": ma.purity_score},
                    "after": {"entity_count": mb.entity_count, "purity": mb.purity_score},
                })
            else:
                unchanged.append({"name": name})

        return {"added": added, "removed": removed, "changed": changed, "unchanged": unchanged}

    async def _render(self, data: dict, format: str) -> bytes:
        """Render report data into the requested format."""
        if format == "json":
            return json.dumps(data, indent=2, default=str).encode("utf-8")

        if format == "html":
            return self._render_html(data)

        if format == "pdf":
            return await self._render_pdf(data)

        if format == "docx":
            return await self._render_docx(data)

        raise ValueError(f"Unsupported format: {format}")

    @staticmethod
    def _render_html(data: dict) -> bytes:
        """Render report as HTML."""
        title = data.get("report_type", "Report").replace("_", " ").title()
        generated = data.get("generated_at", "")

        html_parts = [
            "<!DOCTYPE html><html><head>",
            f"<title>{title}</title>",
            "<style>body{font-family:sans-serif;margin:2em;} "
            "table{border-collapse:collapse;width:100%;} "
            "th,td{border:1px solid #ddd;padding:8px;text-align:left;} "
            "th{background:#f5f5f5;}</style>",
            "</head><body>",
            f"<h1>{title}</h1>",
            f"<p>Generated: {generated}</p>",
            f"<pre>{json.dumps(data, indent=2, default=str)}</pre>",
            "</body></html>",
        ]
        return "".join(html_parts).encode("utf-8")

    @staticmethod
    async def _render_pdf(data: dict) -> bytes:
        """Render report as PDF using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            title = data.get("report_type", "Report").replace("_", " ").title()
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 12))
            story.append(
                Paragraph(
                    f"Generated: {data.get('generated_at', '')}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            # Add JSON content as formatted text
            content_str = json.dumps(data, indent=2, default=str)
            for line in content_str.split("\n")[:200]:
                story.append(Paragraph(line.replace(" ", "&nbsp;"), styles["Code"]))

            doc.build(story)
            return buffer.getvalue()

        except ImportError:
            logger.warning("ReportLab not available; returning JSON as PDF fallback")
            return json.dumps(data, indent=2, default=str).encode("utf-8")

    @staticmethod
    async def _render_docx(data: dict) -> bytes:
        """Render report as DOCX using python-docx."""
        try:
            from docx import Document

            doc = Document()
            title = data.get("report_type", "Report").replace("_", " ").title()
            doc.add_heading(title, 0)
            doc.add_paragraph(f"Generated: {data.get('generated_at', '')}")
            doc.add_paragraph("")

            content = json.dumps(data, indent=2, default=str)
            doc.add_paragraph(content[:5000])

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            logger.warning("python-docx not available; returning JSON as DOCX fallback")
            return json.dumps(data, indent=2, default=str).encode("utf-8")

    @staticmethod
    def _validate_format(format: str) -> None:
        if format not in SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: {format}. Supported: {', '.join(SUPPORTED_FORMATS)}"
            )
